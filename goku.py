from docx import Document
from colorama import init,Fore,Style
import os,subprocess


document  = Document()
def read_code(file:str)->str:
    code = ""
    with open(file,'r') as fl:
        code = fl.read()

    return code


#This function is to only get specifically c files 
#TODO: Lets Make this multilanguage support guys 
def handle_c_files() -> list:
    c_files = []
    cwd = os.getcwd()
    files = os.listdir(cwd)
    for file in files:
        for i in range(0,len(file)):
            if file[i] == '.':
                extension  = file[i:]
                if extension == ".c":
                    c_files.append(file)  
    
    return c_files


def get_file_name(file:str):
    filename = ""
    for  i in range(0,len(file)):
        if  file[i] == ".":
            filename = file[:i]
    return filename

    


def output_generator()->dict:
    current_task = None
    tasks = {}
    current_content = []
    with open("output.gok",'r') as file:
        for line in file:
            line = line.strip()
            if line.upper().startswith("TASK"):
                if current_task:
                    tasks[current_task] = '\n'.join(current_content).strip()
                current_task = line
                
                current_content = []

            elif current_task:
                current_content.append(line)

        if current_content and current_task:
                    tasks[current_task] = '\n'.join(current_content).strip()

    print(f"{Fore.GREEN}The output was successfully read from output.gok{Style.RESET_ALL}")

    return tasks


def details():
    print("Please enter your name and registration number below:")
    name = input("Name:")
    reg_no = input("Registration number:")
    document.add_heading("Assignment Submission",0).add_run().underline = True
    document.add_heading(f"Name:{name.upper()}",1)
    document.add_heading(f"Registration number:{reg_no.upper()}",1)
    document.save("submission.docx")
def file_generator(code:str,id:int,output:str):
    try:
        h1 = document.add_heading(f"Task {id}")
        h1.add_run().underline = True
        code_header = document.add_heading("Code",2)
        code_header.add_run().underline = True
        code_text = document.add_paragraph(code)
        output_header = document.add_heading("Output",2)
        output_header.add_run().underline = True
        output_text = document.add_paragraph(output)
        document.save("submission.docx")
    except Exception as e:
        print(f"File generation process encountered an error {e}")

        
def main():
    init()
    print(f"{Fore.BLUE}{Style.BRIGHT}GOKU TOOL㊗㊗㊗㊗㊗{Style.RESET_ALL}")
    details()
    index = 1
    for file in handle_c_files():
      
        code = read_code(file)
      
       
        output_list = output_generator()
        file_generator(code,index,output_list[F"TASK {index}"])
        index+=1
    if(os.path.isfile("submission.docx")):
        print(f"{Fore.GREEN}The document generation was successful!!!.Look for word.docx in your current working directory{Style.RESET_ALL}")
    else:
        print(f"{Fore.RED}{Style.BRIGHT}For some reason the process was unsuccessful.Please try again")
    


if __name__ == "__main__":
    main()